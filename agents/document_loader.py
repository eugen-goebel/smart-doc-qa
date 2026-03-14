"""
Document Loader — Reads PDF, DOCX, and TXT files and extracts plain text.

Think of this as a "translator" that converts different file formats into
a single format (plain text) that the rest of the pipeline can work with.

Supported formats:
  - .pdf  → uses pypdf to extract text from each page
  - .docx → uses python-docx to read paragraphs
  - .txt  → reads the file directly
"""

import os
from pypdf import PdfReader
from docx import Document
from pydantic import BaseModel, Field


# ---------------------------------------------------------------------------
# Data model — describes what we know about a loaded document
# ---------------------------------------------------------------------------

class LoadedDocument(BaseModel):
    """
    Represents a document that has been read and converted to text.

    Attributes:
        filename:   Original file name (e.g. "report.pdf")
        format:     File type ("pdf", "docx", or "txt")
        text:       The full extracted text content
        page_count: Number of pages (only meaningful for PDFs)
        char_count: Total number of characters in the text
    """
    filename: str = Field(description="Original filename")
    format: str = Field(description="File format: pdf, docx, or txt")
    text: str = Field(description="Full extracted text")
    page_count: int = Field(description="Number of pages (1 for txt/docx)")
    char_count: int = Field(description="Total characters in the text")


# ---------------------------------------------------------------------------
# Supported file extensions
# ---------------------------------------------------------------------------

SUPPORTED_FORMATS = {".pdf", ".docx", ".txt"}


# ---------------------------------------------------------------------------
# Agent
# ---------------------------------------------------------------------------

class DocumentLoader:
    """
    Reads a file from disk and returns its text content.

    Usage:
        loader = DocumentLoader()
        doc = loader.load("path/to/file.pdf")
        print(doc.text)        # the full text
        print(doc.page_count)  # number of pages
    """

    def load(self, filepath: str) -> LoadedDocument:
        """
        Load a document and extract its text.

        Args:
            filepath: Path to a .pdf, .docx, or .txt file

        Returns:
            LoadedDocument with the extracted text and metadata

        Raises:
            FileNotFoundError: If the file does not exist
            ValueError: If the file format is not supported
        """
        # --- Step 1: Check that the file exists ---
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        # --- Step 2: Determine the file format ---
        ext = os.path.splitext(filepath)[1].lower()
        if ext not in SUPPORTED_FORMATS:
            raise ValueError(
                f"Unsupported format: '{ext}'. "
                f"Use one of: {', '.join(sorted(SUPPORTED_FORMATS))}"
            )

        # --- Step 3: Read the file using the appropriate method ---
        if ext == ".pdf":
            text, page_count = self._read_pdf(filepath)
        elif ext == ".docx":
            text, page_count = self._read_docx(filepath)
        else:
            text, page_count = self._read_txt(filepath)

        # --- Step 4: Package the result ---
        return LoadedDocument(
            filename=os.path.basename(filepath),
            format=ext.lstrip("."),
            text=text,
            page_count=page_count,
            char_count=len(text),
        )

    def _read_pdf(self, filepath: str) -> tuple[str, int]:
        """
        Extract text from a PDF file.

        pypdf reads each page individually, so we join them with
        double newlines to preserve page boundaries.
        """
        reader = PdfReader(filepath)
        pages = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                pages.append(page_text.strip())

        return "\n\n".join(pages), len(reader.pages)

    def _read_docx(self, filepath: str) -> tuple[str, int]:
        """
        Extract text from a Word (.docx) file.

        python-docx reads paragraphs. We join non-empty paragraphs
        with newlines.
        """
        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs), 1

    def _read_txt(self, filepath: str) -> tuple[str, int]:
        """Read a plain text file."""
        with open(filepath, "r", encoding="utf-8") as f:
            text = f.read()
        return text, 1
